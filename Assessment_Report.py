
# Main function to create a report of a pentest

# Python Modules
import os, docx, logging, math
from docx.enum.text import WD_ALIGN_PARAGRAPH
# My Modules
import DB_Connect, Retrieve_CVSS, Sort_IP, Print_Banner, Main_Menu

#################################################
#                   COLORS                      #
#################################################
note = "\033[0;0;33m-\033[0m"
warn = "\033[0;0;31m!\033[0m"
info = "\033[0;0;36mi\033[0m"
question = "\033[0;0;37m?\033[0m"
debug = "\033[0;0;31m[DEBUG]\033[0m"


def assessment_vulns(assessment, crosstable):
    """Builds a list of the assessment vulnerabilities"""

    vulns = []
    plugins = ""

    #Import data from gauntlet db for the selected crosstable
    hosts = DB_Connect.db_query("select * from cross_data_nva WHERE table_id = '%s'" % (crosstable), assessment)

    # Vureto Vuln ID is in spot 2 of the tuple returned by i in the following loop
    Vureto_Vuln_ID_Location = 2

    for i in hosts:
        Vureto_Vuln_ID = i[Vureto_Vuln_ID_Location]
        if ((Vureto_Vuln_ID.startswith('Nessus') or Vureto_Vuln_ID.startswith('Netsparker') or Vureto_Vuln_ID.startswith('Acunetix')
                or Vureto_Vuln_ID.startswith('BurpSuite')) and Vureto_Vuln_ID not in plugins):
            plugins += Vureto_Vuln_ID
        else:
            vulns.append(Vureto_Vuln_ID)
    if plugins != "":
        pass
    return vulns

def get_vulns(vuln_IDs, assessment, crosstable, cli_args):
    """Build dictionary containing the assessment vulnerabilities and their associated information"""

    logging.info('Entered get_vulns function...')
    vulns = {}
    plugins = ""
    tools = ['Nessus', 'Netsparker', 'Acunetix', 'BurpSuite', 'Nmap', 'Nikto', 'dirb']  # names of tools to ignore

    # Connect to the database
    db = DB_Connect.db_connect('GauntletData', cli_args)
    DB_Connect.ignore_mysql_warnings()

    # Variable for progress bar
    countvulnVar = 0

    # vuln_IDs is an array of vuln id's created in assessment_vulns function
    for vuln_id in vuln_IDs:

        # Create a progress bar
        countvulnVar = countvulnVar + 1.0
        progress = str(countvulnVar / len(vuln_IDs) * 100)
        print "\r\t[" + note + "]Querying Database For Vulnerability Information:", progress[:5] + "%",

        #TODO to remove "Nessus 1111" entries

        if vuln_id.split()[0] in tools:
            if vuln_id not in plugins and vuln_id is "":
                plugins += "\t["+warn+"]" + vuln_id + " plugin needs to be added to your Gauntlet database"
            elif vuln_id not in plugins:
                plugins += "\n\t["+warn+"]" + vuln_id + " plugin needs to be added to your Gauntlet database"

        # Get vulnerability verbiage
        else:
            q = """select title, description, solution, report_id from vulns WHERE gnaat_id=%s""" % (vuln_id)
            temp = DB_Connect.db_interact_fetchone(db, q)

            if temp[3] is not None:
                vulns[vuln_id] = {'vuln_id': vuln_id, 'vuln_title': temp[0], 'vuln_desc': temp[1], 'vuln_sol': temp[2], 'vuln_report_id': int(temp[3])}
            else:
                vulns[vuln_id] = {'vuln_id': vuln_id, 'vuln_title': temp[0], 'vuln_desc': temp[1], 'vuln_sol': temp[2], 'vuln_report_id': temp[3]}

    # Connect to the database
    db2 = DB_Connect.db_connect('gauntlet_' + assessment, cli_args)

    #Add all hosts with the associated vulnerability to the rpt dictionary
    print ""
    countvulnVar = 0

    #Determine if cross_data_nva has a 'port' column
    portColumn = False

    columns = DB_Connect.db_interact_fetchall(db2, """SHOW COLUMNS FROM cross_data_nva FROM gauntlet_%s;""" % (assessment))

    for c in columns:
        if c[0] == 'port':
            portColumn = True

    for vuln_id in vuln_IDs:

        # Create progress bar
        countvulnVar = countvulnVar + 1.0
        progress = str(countvulnVar / len(vuln_IDs) * 100)
        print "\r\t[" + note + "]Querying Database For Affected Hosts:", progress[:5] + "%",

        # Ignore id's in tool list
        if vuln_id.split()[0] in tools:
            pass
        else:
            temp_db_cursor = db2.cursor()
            if cli_args.ports and portColumn:
                temp_db_cursor.execute(
                    """SELECT host, port, protocol FROM cross_data_nva WHERE cross_data_nva.table_id =%s
                    AND vuln_id=%s AND (s1='Y' or s2='Y' or s3='Y' or s4='Y' or s5='Y')""",
                    (crosstable,vuln_id))
            else:
                temp_db_cursor.execute(
                    """SELECT host FROM cross_data_nva WHERE cross_data_nva.table_id =%s AND vuln_id=%s
                    AND (s1='Y' or s2='Y' or s3='Y' or s4='Y' or s5='Y')""",
                    (crosstable, vuln_id))
            temp_cursor_returned_info = temp_db_cursor.fetchall()
            temp_db_cursor.close()

            vulns[vuln_id].update({'vuln_hosts': temp_cursor_returned_info})

    #Determine the rank of the vulnerability
    print ""
    countvulnVar = 0
    for vuln_id in vuln_IDs:

        # Create progress bar
        countvulnVar = countvulnVar + 1.0
        progress = str(countvulnVar / len(vuln_IDs) * 100)
        print "\r\t[" + note + "]Querying Database For Vulnerability Rank:", progress[:5] + "%",

        # Ignore ids in tool list
        if vuln_id.split()[0] in tools:
            pass
        # Select severity values from cross_tables
        else:
            severity_tuple = DB_Connect.db_query("""SELECT s1, s2, s3, s4, s5 FROM cross_data_nva WHERE table_id ='%s' 
            AND vuln_id = %s;""" % (crosstable, vuln_id), assessment)
            severities = []
            for value in severity_tuple:
                if value[0] is 'Y' and ('Critical' not in severities):
                    severities.append('Critical')
                elif value[1] is 'Y' and ('High' not in severities):
                    severities.append('High')
                elif value[2] is 'Y' and ('Medium' not in severities):
                    severities.append('Medium')
                elif value[3] is 'Y' and ('Low' not in severities):
                    severities.append('Low')
                elif value[4] is 'Y' and ('Informational' not in severities):
                    severities.append('Informational')
                else:
                    if None in severities:
                        pass
                    else:
                        severities.append(None)

            # Update vuln_rating to the highest discovered severity
            if 'Critical' in severities:
                vulns[vuln_id].update({'vuln_rating': 'Critical'})
            elif 'High' in severities:
                vulns[vuln_id].update({'vuln_rating': 'High'})
            elif 'Medium' in severities:
                vulns[vuln_id].update({'vuln_rating': 'Medium'})
            elif 'Low' in severities:
                vulns[vuln_id].update({'vuln_rating': 'Low'})
            elif 'Informational' in severities:
                vulns[vuln_id].update({'vuln_rating': 'Informational'})
            else:
                vulns[vuln_id].update({'vuln_rating': None})

    if plugins != "":
        print "%s", (plugins),
    print ""
    DB_Connect.dont_ignore_mysql_warnings()
    return vulns

def assessment_report(vulnerability_id_info_mapping_list, cli_args):
    """Builds a unique list of Report IDs for the selected assessment and crosstable"""

    logging.info('Entered assessment_report function')
    vuln_report_id_list = []
    for vulnerability_id in vulnerability_id_info_mapping_list:
        if cli_args.debug:
            logging.info('Vulnerability: ' + vulnerability_id + " - " + vulnerability_id_info_mapping_list[vulnerability_id]['vuln_title'])
        if vulnerability_id_info_mapping_list[vulnerability_id]['vuln_report_id'] is not None:
            vuln_report_id_list.append(vulnerability_id_info_mapping_list[vulnerability_id]['vuln_report_id'])
        else:
            pass
    logging.info('Leaving assessment_report_function')
    return set(vuln_report_id_list)


def get_report(report_record_IDs, vuln_ID_info_mapping, cli_args):
    """Build a dictionary containing all of the reporting information"""
    logging.info("Entering get_report function")

    rpt = {}

    # change to GauntletData after dev/or vureto for dev
    db = DB_Connect.db_connect('GauntletData', cli_args)

    for report_record_ID in report_record_IDs:
        query = """select title, identification, explanation, impact, recommendation from report
                         WHERE report_id=%s""" % (report_record_ID)
        report_record_dictionary = DB_Connect.db_interact_fetchone(db, query)

        rpt[report_record_ID] = {'report_id': report_record_ID, 'report_title': report_record_dictionary[0],
                                 'report_identification': report_record_dictionary[1],
                                 'report_explanation': report_record_dictionary[2], 'report_impact': report_record_dictionary[3],
                                 'report_recommendation': report_record_dictionary[4]}

    # Add all vulnerabilities with this report ID to the dictionary
    for vuln_ID in vuln_ID_info_mapping:
        if vuln_ID_info_mapping[vuln_ID]['vuln_report_id'] is not None:
            if 'vulns' in rpt[vuln_ID_info_mapping[vuln_ID]['vuln_report_id']]:
                rpt[vuln_ID_info_mapping[vuln_ID]['vuln_report_id']]['vulns'][
                    vuln_ID_info_mapping[vuln_ID]['vuln_id']] = vuln_ID_info_mapping[vuln_ID]
            else:
                rpt[vuln_ID_info_mapping[vuln_ID]['vuln_report_id']]['vulns'] = {
                    vuln_ID_info_mapping[vuln_ID]['vuln_id']: vuln_ID_info_mapping[vuln_ID]}
        else:
            pass

    # Determine the highest severity level and set it for the reporting record
    for record in rpt:
        r = []
        for vulnerability in rpt[record]['vulns']:
            r.append(rpt[record]['vulns'][vulnerability]['vuln_rating'])
        if 'Critical' in r:
            rpt[record]['report_rating'] = 'Critical'
            continue
        elif 'High' in r:
            rpt[record]['report_rating'] = 'High'
            continue
        elif 'Medium' in r:
            rpt[record]['report_rating'] = 'Medium'
            continue
        elif 'Low' in r:
            rpt[record]['report_rating'] = 'Low'
            continue
        elif 'Informational' in r:
            rpt[record]['report_rating'] = 'Informational'
            continue
        else:
            rpt[record]['report_rating'] = None
    logging.info('Leaving get_report Function')
    return rpt



def generate_vuln_list(report_docx, assessment, vuln_dictionary_data, cli_args):
    """Build the bullet list of vulnerabilities used in the executive summary"""

    logging.info("Entering the generate_vuln_list function")
    DB_Connect.ignore_mysql_warnings()
    engagement = DB_Connect.db_query("""SELECT value FROM gauntlet_%s.engagement_details WHERE engagement_details.key = 
                            'Engagement Task 1'""" % assessment, assessment)
    if engagement:
        report_docx.add_heading(str(engagement[0][0]) + ' NVA/PT')
    else:
        report_docx.add_heading('NVA/PT')

    def writeBullet(s, h):
        n = s.find('[n]') + 3  # Find '[n]' and add three to account for the length of '[n]'
        s = s.replace(s[0:n], int_to_string(h))
        s = s.rstrip('\n')
        s = s.rstrip('\t')
        s = s.rstrip()
        s = s.rstrip('.')
        s = s[0:1].upper() + s[1:]
        report_docx.add_paragraph(s, style='List Bullet')

    logging.info("Writing bullets based on criticality in the generate_vuln_list function")
    for i in vuln_dictionary_data:
        if cli_args.debug:
            print "[" + info + "]%s" % vuln_dictionary_data[i]

        # Check to see if is a multi vuln report item
        if len(vuln_dictionary_data[i]['vulns']) > 1:
            h = 0
            for j in vuln_dictionary_data[i]['vulns']:
                h += len(vuln_dictionary_data[i]['vulns'][j]['vuln_hosts'])
            if cli_args.sC and vuln_dictionary_data[i]['report_rating'] == 'Critical':
                writeBullet(vuln_dictionary_data[i]['report_identification'], h)
            elif cli_args.sH and vuln_dictionary_data[i]['report_rating'] == 'High':
                writeBullet(vuln_dictionary_data[i]['report_identification'], h)
            elif cli_args.sM and vuln_dictionary_data[i]['report_rating'] == 'Medium':
                writeBullet(vuln_dictionary_data[i]['report_identification'], h)
            elif cli_args.sL and vuln_dictionary_data[i]['report_rating'] == 'Low':
                writeBullet(vuln_dictionary_data[i]['report_identification'], h)
            elif cli_args.sI and vuln_dictionary_data[i]['report_rating'] == 'Informational':
                writeBullet(vuln_dictionary_data[i]['report_identification'], h)
            elif vuln_dictionary_data[i]['report_rating'] is None:
                print "\t[" + note + "]" + vuln_dictionary_data[i]['report_title'] + " has no affected hosts"
            else:
                pass

        else:
            for j in vuln_dictionary_data[i]['vulns']:
                if cli_args.sC and vuln_dictionary_data[i]['report_rating'] == 'Critical':
                    writeBullet(vuln_dictionary_data[i]['report_identification'], len(vuln_dictionary_data[i]['vulns'][j]['vuln_hosts']))
                elif cli_args.sH and vuln_dictionary_data[i]['report_rating'] == 'High':
                    writeBullet(vuln_dictionary_data[i]['report_identification'], len(vuln_dictionary_data[i]['vulns'][j]['vuln_hosts']))
                elif cli_args.sM and vuln_dictionary_data[i]['report_rating'] == 'Medium':
                    writeBullet(vuln_dictionary_data[i]['report_identification'], len(vuln_dictionary_data[i]['vulns'][j]['vuln_hosts']))
                elif cli_args.sL and vuln_dictionary_data[i]['report_rating'] == 'Low':
                    writeBullet(vuln_dictionary_data[i]['report_identification'], len(vuln_dictionary_data[i]['vulns'][j]['vuln_hosts']))
                elif cli_args.sI and vuln_dictionary_data[i]['report_rating'] == 'Informational':
                    writeBullet(vuln_dictionary_data[i]['report_identification'], len(vuln_dictionary_data[i]['vulns'][j]['vuln_hosts']))
                elif vuln_dictionary_data[i]['report_rating'] is None:
                    print "\t[" + note + "]" + vuln_dictionary_data[i]['report_title'] + " has no affected hosts"
                else:
                    pass

    return report_docx


def int_to_string(i):
    """Converts an integer to its spelled out version; Used in reporting narratives"""

    s = {0: "",  # Dictionary of integers to spelled out words
         1: "one",
         2: "two",
         3: "three",
         4: "four",
         5: "five",
         6: "six",
         7: "seven",
         8: "eight",
         9: "nine",
         10: "ten",
         11: "eleven",
         12: "twelve",
         13: "thirteen",
         14: "fourteen",
         15: "fifteen",
         16: "sixteen",
         17: "seventeen",
         18: "eighteen",
         19: "nineteen",
         20: "twenty",
         30: "thirty",
         40: "forty",
         50: "fifty",
         60: "sixty",
         70: "seventy",
         80: "eighty",
         90: "ninety"
         }

    # break i into ones, tens, hundreds, thousands and then build string
    # ONES
    if len(str(i)) is 1 and str(i) is "1":  # Spelling for host as opposed to hosts
        return s[i] + " (" + str(i) + ") host"
    elif len(str(i)) is 1 and str(i) is not "1":  # Spelling for single digit hosts
        return s[i] + " (" + str(i) + ") hosts"
    # TENS
    elif len(str(i)) is 2 and str(i).startswith("1"):  # To grab spelling for 11 through 19
        return s[i] + " (" + str(i) + ") hosts"
    elif len(str(i)) is 2 and not str(i).startswith("1") and not str(i).endswith(
            "0"):  # To grab spelling for 20 through 99 where the number doesn't end in 0
        return s[(i / 10) * 10] + "-" + s[i - ((i / 10) * 10)] + " (" + str(i) + ") hosts"
    elif len(str(i)) is 2 and not str(i).startswith("1") and str(i).endswith(
            "0"):  # To grab spelling for 20 through 99 where the number doesn't end in 0
        return s[(i / 10) * 10] + s[i - ((i / 10) * 10)] + " (" + str(i) + ") hosts"
    # HUNDREDS
    elif len(str(i)) is 3 and "0" not in str(i):  # to grab spelling for 100's where the number doesn't have a zero
        return s[(i / 100)] + "-hundred " + s[((i - ((i / 100) * 100)) / 10) * 10] + "-" + s[
            i - (((i / 100) * 100) + ((i - ((i / 100) * 100)) / 10) * 10)] + " (" + str(i) + ") hosts"
    elif len(str(i)) is 3 and not str(i).endswith("0") and "0" in str(
            i):  # to grab spelling for 100's where the 10s place is 0
        return s[(i / 100)] + "-hundred" + s[((i - ((i / 100) * 100)) / 10) * 10] + " " + s[
            i - (((i / 100) * 100) + ((i - ((i / 100) * 100)) / 10) * 10)] + " (" + str(i) + ") hosts"
    elif len(str(i)) is 3 and str(i).endswith("0"):  # to grab spelling for 100's where the number ends in 0
        return s[(i / 100)] + "-hundred" + s[((i - ((i / 100) * 100)) / 10) * 10] + s[
            i - (((i / 100) * 100) + ((i - ((i / 100) * 100)) / 10) * 10)] + " (" + str(i) + ") hosts"
    else:
        return "ERROR, was not able to return the number of host(s)"

def write_multi_vul(rpt, report_docx, cli_args):
    """Write report data for grouped or multi vulnerabilities"""

    total_hosts = 0
    for vulnerability in rpt['vulns']:
        if cli_args.sC and rpt['vulns'][vulnerability]['vuln_rating'] == 'Critical':
            total_hosts += len(rpt['vulns'][vulnerability]['vuln_hosts'])
        elif cli_args.sH and rpt['vulns'][vulnerability]['vuln_rating'] == 'High':
            total_hosts += len(rpt['vulns'][vulnerability]['vuln_hosts'])
        elif cli_args.sM and rpt['vulns'][vulnerability]['vuln_rating'] == 'Medium':
            total_hosts += len(rpt['vulns'][vulnerability]['vuln_hosts'])
        elif cli_args.sL and rpt['vulns'][vulnerability]['vuln_rating'] == 'Low':
            total_hosts += len(rpt['vulns'][vulnerability]['vuln_hosts'])
        elif cli_args.sI and rpt['vulns'][vulnerability]['vuln_rating'] == 'Informational':
            total_hosts += len(rpt['vulns'][vulnerability]['vuln_hosts'])

    report_docx.add_heading(rpt['report_title'] + " (" + rpt['report_rating'] + ")", level=4)
    p = rpt['report_identification'].replace("[n]", int_to_string(total_hosts))
    p = p.rstrip('\n')
    p = p.rstrip('\t')
    p = p.rstrip()

    if p.endswith(" ") or rpt['report_explanation'].startswith(" "):
        p += rpt['report_explanation']
    else:
        p += (" " + rpt['report_explanation'])
    if p.endswith(" ") or rpt['report_impact'].startswith(" "):
        p += rpt['report_impact']
    else:
        p += (" " + rpt['report_impact'])

    p = report_docx.add_paragraph(p, style='Normal')
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p = report_docx.add_paragraph(rpt['report_recommendation'], style='Normal')
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    table = None
    hdr_cells = None

    if cli_args.cvss:
        table = report_docx.add_table(rows=1, cols=4)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Severity'
        hdr_cells[1].text = 'CVSS'
        hdr_cells[2].text = 'Vulnerability'
        hdr_cells[3].text = 'Affected Host(s)'
    else:
        table = report_docx.add_table(rows=1, cols=3)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Severity'
        hdr_cells[1].text = 'Vulnerability'
        hdr_cells[2].text = 'Affected Host(s)'

    table.style = 'Medium Grid 1 Accent 1'

    def writeRow(rating, title, hosts, cvss_score=None):
        row_cells = table.add_row().cells
        if cvss_score is not None:
            row_cells[0].text = rating
            row_cells[1].text = cvss_score
            row_cells[2].text = title
            row_cells[3].text = hosts
        else:
            row_cells[0].text = rating
            row_cells[1].text = title
            row_cells[2].text = hosts

    for vulnerability in rpt['vulns']:
        if cli_args.sC and rpt['vulns'][vulnerability]['vuln_rating'] == 'Critical':
            cvss = Retrieve_CVSS.get_cvss(vulnerability)
            if cli_args.cvss:
                writeRow(rpt['vulns'][vulnerability]['vuln_rating'], rpt['vulns'][vulnerability]['vuln_title'],
                         str(len(rpt['vulns'][vulnerability]['vuln_hosts'])), "%0.1f" % cvss.base_score)
            else:
                writeRow(rpt['vulns'][vulnerability]['vuln_rating'], rpt['vulns'][vulnerability]['vuln_title'],
                         str(len(rpt['vulns'][vulnerability]['vuln_hosts'])))
    for vulnerability in rpt['vulns']:
        if cli_args.sH and rpt['vulns'][vulnerability]['vuln_rating'] == 'High':
            cvss = Retrieve_CVSS.get_cvss(vulnerability)
            if cli_args.cvss:
                writeRow(rpt['vulns'][vulnerability]['vuln_rating'], rpt['vulns'][vulnerability]['vuln_title'],
                         str(len(rpt['vulns'][vulnerability]['vuln_hosts'])), "%0.1f" % cvss.base_score)
            else:
                writeRow(rpt['vulns'][vulnerability]['vuln_rating'], rpt['vulns'][vulnerability]['vuln_title'],
                         str(len(rpt['vulns'][vulnerability]['vuln_hosts'])))
    for vulnerability in rpt['vulns']:
        if cli_args.sM and rpt['vulns'][vulnerability]['vuln_rating'] == 'Medium':
            cvss = Retrieve_CVSS.get_cvss(vulnerability)
            if cli_args.cvss:
                writeRow(rpt['vulns'][vulnerability]['vuln_rating'], rpt['vulns'][vulnerability]['vuln_title'],
                         str(len(rpt['vulns'][vulnerability]['vuln_hosts'])), "%0.1f" % cvss.base_score)
            else:
                writeRow(rpt['vulns'][vulnerability]['vuln_rating'], rpt['vulns'][vulnerability]['vuln_title'],
                         str(len(rpt['vulns'][vulnerability]['vuln_hosts'])))
    for vulnerability in rpt['vulns']:
        if cli_args.sL and rpt['vulns'][vulnerability]['vuln_rating'] == 'Low':
            cvss = Retrieve_CVSS.get_cvss(vulnerability)
            if cli_args.cvss:
                writeRow(rpt['vulns'][vulnerability]['vuln_rating'], rpt['vulns'][vulnerability]['vuln_title'],
                         str(len(rpt['vulns'][vulnerability]['vuln_hosts'])), "%0.1f" % cvss.base_score)
            else:
                writeRow(rpt['vulns'][vulnerability]['vuln_rating'], rpt['vulns'][vulnerability]['vuln_title'],
                         str(len(rpt['vulns'][vulnerability]['vuln_hosts'])))
    for vulnerability in rpt['vulns']:
        if cli_args.sI and rpt['vulns'][vulnerability]['vuln_rating'] == 'Informational':
            cvss = Retrieve_CVSS.get_cvss(vulnerability)
            if cli_args.cvss:
                writeRow(rpt['vulns'][vulnerability]['vuln_rating'], rpt['vulns'][vulnerability]['vuln_title'],
                         str(len(rpt['vulns'][vulnerability]['vuln_hosts'])), "%0.1f" % cvss.base_score)
            else:
                writeRow(rpt['vulns'][vulnerability]['vuln_rating'], rpt['vulns'][vulnerability]['vuln_title'],
                         str(len(rpt['vulns'][vulnerability]['vuln_hosts'])))

    return report_docx

def write_single_vul(rpt, report_docx, cli_args):
    """Write the single vulnerability paragraph"""

    cvss = None
    if len(rpt['vulns']) is 1:
        for rows in rpt['vulns']:
            cvss = Retrieve_CVSS.get_cvss(rows)

    if cli_args.cvss:
        report_docx.add_heading("%s (CVSS: %0.1f - %s)" % (rpt['report_title'], cvss.base_score, rpt['report_rating']),
                                level=4)
    else:
        report_docx.add_heading("%s (%s)" % (rpt['report_title'], rpt['report_rating']), level=4)

    for vulnerability in rpt['vulns']:
        vulnerableHosts = []
        for host in set(rpt['vulns'][vulnerability]['vuln_hosts']):  # Build Single Dimensional List
            vulnerableHosts.append(host)
        sortedHosts = Sort_IP.ip_sort_tuple(vulnerableHosts)  # Create a unique & sorted list of hosts (avoids duplicate hosts)
        hosts = []

        #TODO Make this understandable lol
        for h in sortedHosts:
            for h2 in vulnerableHosts:
                if h2[0] == h:
                    if cli_args.ports:
                        if len(h2) == 3:
                            if h2[1] != '0' and h2[2] != 'icmp':
                                host_info = h2[0] + ":" + h2[1] + "/" + h2[2]
                                if host_info not in hosts:
                                    hosts.append(host_info)
                            else:
                                if h2[0] not in hosts:
                                    hosts.append(h2[0])
                        else:
                            if h2[0] not in hosts:
                                hosts.append(h2[0])
                    else:
                        if h2[0] not in hosts:
                            hosts.append(h2[0])

        if len(hosts) == 1:         # If there is just one host
            p = rpt['report_identification'].replace("[n]", int_to_string(len(hosts))+ ' ('+hosts[0]+')')
        elif len(hosts) == 2:       # If there are two hosts
            p = rpt['report_identification'].replace("[n]", int_to_string(len(hosts))+ ' ('+hosts[0]+' and '+hosts[1]+')')
        elif len(hosts) >= 2 and len(hosts) <=5:  # If there are more than two but less than five hosts
            host_list = ""
            for h in hosts:
                if h is hosts[len(hosts)-1]:  # Check to see if this is the last item in the list
                    host_list += "and " + h + ") "
                else:
                    host_list += h + ", "
            p = rpt['report_identification'].replace("[n]", int_to_string(len(hosts)) + ' ('+host_list)
        elif len(hosts) >= 6:  # If there are six or more hosts
            p = rpt['report_identification'].replace("[n]", int_to_string(len(hosts)) + '(refer to TABLE X)')

        if p.endswith(" ") or rpt['report_explanation'].startswith(" "):
            p += rpt['report_explanation']
        else:
            p += (" " + rpt['report_explanation'])
        if p.endswith(" ") or rpt['report_impact'].startswith(" "):
            p += rpt['report_impact']
        else:
            p += (" " + rpt['report_impact'])
        p = report_docx.add_paragraph(p, style='Normal')
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p = report_docx.add_paragraph(rpt['report_recommendation'], style='Normal')
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        if cli_args.debug:
            print debug + "Vulnerable Hosts: %s" % hosts
            raw_input(debug + "Press any key to continue")
        if len(hosts) >= 6:     # Draw the table
            columns = 4  # number of desired columns
            rows = int(math.ceil((len(hosts) / float(4))))  # Determine number of rows for table using a max of 4 columns
            hosts_table = report_docx.add_table(rows=rows, cols=columns)
            hosts_table.style = 'Medium Grid 1 Accent 1'
            host_counter = 0   # number of hosts
            x = 0   # row indices
            y = 0   # column indices
            while host_counter < len(hosts):
                if (y / float(columns)) == 1:  # Determine if we need to start putting data on a new row
                    y = 0   # reset column indices since max number of columns reached
                    x += 1
                hosts_table.cell(x, y).text = hosts[host_counter]
                host_counter += 1
                y += 1  # Add one to up the column data is put in
            if len(hosts)/float(columns) != 1.000:  # Add "---" for empty spots in table
                d = columns * (x+1)
                while d > len(hosts):
                    hosts_table.cell(x, y).text = "---"
                    d -= 1
                    y += 1

    return report_docx


def write_all_vuln(vuln, report_docx, cli_args):

    print "["+note+"]Writing list of all vulnerabilities to the report: "
    report_docx.add_page_break()
    report_docx.add_heading("List of Assessment Vulnerabilities", 1)
    for i in vuln:
        if cli_args.sC and vuln[i]['vuln_rating'] is 'Critical':
            print "\t["+info+"]"+ vuln[i]['vuln_title'], "(" + vuln[i]['vuln_rating'] + ")"
            report_docx.add_heading(vuln[i]['vuln_title'] + " (" + vuln[i]['vuln_rating'] + ")", 3)
            if cli_args.all_verb:
                report_docx.add_paragraph(vuln[i]['vuln_desc'], style='Body Text')
                report_docx.add_paragraph(vuln[i]['vuln_sol'], style='Body Text')
    for i in vuln:
        if cli_args.sH and vuln[i]['vuln_rating'] is 'High':
            print "\t["+info+"]"+ vuln[i]['vuln_title'], "(" + vuln[i]['vuln_rating'] + ")"
            report_docx.add_heading(vuln[i]['vuln_title'] + " (" + vuln[i]['vuln_rating'] + ")", 3)
            if cli_args.all_verb:
                report_docx.add_paragraph(vuln[i]['vuln_desc'], style='Body Text')
                report_docx.add_paragraph(vuln[i]['vuln_sol'], style='Body Text')
    for i in vuln:
        if cli_args.sM and vuln[i]['vuln_rating'] is 'Medium':
            print "\t["+info+"]"+ vuln[i]['vuln_title'], "(" + vuln[i]['vuln_rating'] + ")"
            report_docx.add_heading(vuln[i]['vuln_title'] + " (" + vuln[i]['vuln_rating'] + ")", 3)
            if cli_args.all_verb:
                report_docx.add_paragraph(vuln[i]['vuln_desc'], style='Body Text')
                report_docx.add_paragraph(vuln[i]['vuln_sol'], style='Body Text')
    for i in vuln:
        if cli_args.sL and vuln[i]['vuln_rating'] is 'Low':
            print "\t["+info+"]"+ vuln[i]['vuln_title'], "(" + vuln[i]['vuln_rating'] + ")"
            report_docx.add_heading(vuln[i]['vuln_title'] + " (" + vuln[i]['vuln_rating'] + ")", 3)
            if cli_args.all_verb:
                report_docx.add_paragraph(vuln[i]['vuln_desc'], style='Body Text')
                report_docx.add_paragraph(vuln[i]['vuln_sol'], style='Body Text')
    for i in vuln:
        if cli_args.sI and vuln[i]['vuln_rating'] is 'Informational':
            print "\t["+info+"]"+ vuln[i]['vuln_title'], "(" + vuln[i]['vuln_rating'] + ")"
            report_docx.add_heading(vuln[i]['vuln_title'] + " (" + vuln[i]['vuln_rating'] + ")", 3)
            if cli_args.all_verb:
                report_docx.add_paragraph(vuln[i]['vuln_desc'], style='Body Text')
                report_docx.add_paragraph(vuln[i]['vuln_sol'], style='Body Text')

    return report_docx

def generate_hosts_table(file, ass):
    """Build a list of assessment interesting hosts; hosts with atleast one TCP or UDP port open."""

    hosts = {}
    file.add_page_break()
    logging.info('Entering the generate_hosts_table function')
    print "["+note+"]Generating Interesting Hosts Table"
    # Build dictionary of host IDs and IPs from gauntlet's 'hosts' table
    engagement = DB_Connect.db_query("""SELECT value FROM gauntlet_%s.engagement_details WHERE engagement_details.key = 'Engagement Task 1'""" % (ass), ass)
    if engagement:
        if 'Internal' in engagement[0][0]:
            temp = DB_Connect.db_query("""SELECT host_id, ip_address, machine_name from hosts""", ass)
        else:
            temp = DB_Connect.db_query("""SELECT host_id, ip_address, fqdn from hosts""", ass)
    else:
        temp = DB_Connect.db_query("""SELECT host_id, ip_address, machine_name from hosts""", ass)
    temp2 = DB_Connect.db_query("""SELECT host_id, port, protocol from ports""", ass)
    for i in temp:
        tcp = []
        udp = []
        for j in temp2:
            if (j[0] == i[0]) and (j[1] != "0"):
                if j[2] == 'tcp':
                    tcp.append(j[1])
                elif j[2] == 'udp':
                    udp.append(j[1])
                else:
                    pass
            else:
                pass
        hosts[i[0]] = {'IP': i[1], 'Name': i[2], 'TCP': tcp, 'UDP': udp}
    x = 0  # Number of interesting hosts counter
    for host in hosts:
        if (len(hosts[host]['TCP']) > 0) or (len(hosts[host]['UDP']) > 0):
            x += 1
        else:
            pass
    logging.info("["+info+"]"+str(x) + " Interesting Hosts")
    print "\t["+info+"]"+str(x) + " Interesting Hosts"
    file.add_heading(str(x) + ' Interesting Host(s) List')
    table = file.add_table(rows=1, cols=4)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'IP Address'
    hdr_cells[1].text = 'Hostname'
    hdr_cells[2].text = 'Open TCP Port(s)'
    hdr_cells[3].text = 'Open UDP Port(s)'
    table.style = 'Medium Grid 1 Accent 1'

    # Build a list of sorted IPs
    sorted_hosts = []
    for host in hosts:
        sorted_hosts.append(hosts[host]['IP'])
    sorted_hosts = Sort_IP.ip_sort_list(sorted_hosts)

    for ip in sorted_hosts:
        for k in hosts:
            if hosts[k]['IP'] == ip:
                if (len(hosts[k]['TCP']) > 0) or (len(hosts[k]['UDP']) > 0):
                    x += 1
                    row_cells = table.add_row().cells
                    row_cells[0].text = hosts[k]['IP']
                    if len(hosts[k]['Name']) > 0:
                        row_cells[1].text = hosts[k]['Name']
                    else:
                        row_cells[1].text = "---"
                    if len(hosts[k]['TCP']) > 0:
                        row_cells[2].text = str(hosts[k]['TCP']).lstrip('[').rstrip(']').replace("'", "")
                    else:
                        row_cells[2].text = "---"
                    if len(hosts[k]['UDP']) > 0:
                        row_cells[3].text = str(hosts[k]['UDP']).lstrip('[').rstrip(']').replace("'", "")
                    else:
                        row_cells[3].text = "---"
                else:
                    pass

    return file

def get_path():
    """Prompt the user to enter a directory path"""

    output_path = None
    while output_path is None:
        print "["+question+"]Please enter the directory where you would like the file saved?"
        output_path = raw_input()
        if os.path.isdir(os.path.expanduser(output_path)):
            pass
        else:
            os.system('clear')
            Print_Banner.print_banner()
            print "["+warn+"]" + str(output_path) + " is not valid, please try again: "
            output_path = None
    return os.path.expanduser(output_path)

def save_report(file, ass):
    """Save the generated assessment report"""
    out_dir = get_path()
    guinevere_file = os.path.join(out_dir, "Guinevere_"+ass+".docx")
    file.save(guinevere_file)
    print "["+warn+"]Report saved to: " + guinevere_file
    raw_input("["+question+"]Press enter to continue...")
    Main_Menu.print_main_menu()


"""The main function for automatically generating an assessment report"""
def generate_assessment_report(cli_args):


    logging.info('Entering the generate_assessment_report function')
    os.system('clear')
    Print_Banner.print_banner()
    print "Retrieving available assessments..."

    # Database containing the options of crosstables
    assessment = DB_Connect.get_assessment("the assessment to create a report for")
    Print_Banner.print_banner()

    # Choose crosstable from the options listed from get_assessment()
    crosstable = DB_Connect.get_crosstable(assessment)

    vulnerability_ID_list = assessment_vulns(assessment, crosstable)
    os.system('clear')
    Print_Banner.print_banner()

    print "["+note+"]Building list of found vulnerabilities for " + assessment + " Crosstable " + crosstable + "..."
    vuln_ID_to_info_mapping = get_vulns(vulnerability_ID_list, assessment, crosstable, cli_args)

    ## FOLLOW vuln_id_to_info_mapping - Contains the severity rating, needs to be reset after each crosstable

    print "["+note+"]Generating report for the following vulnerabilities:"

    # Unique identifier for reporting record, allows access to verbage
    report_record_ID_set = assessment_report(vuln_ID_to_info_mapping, cli_args)

    # Returns huge dictionary of vulnerabilities and associated data
    assessment_db = get_report(report_record_ID_set, vuln_ID_to_info_mapping, cli_args)

    # Make the word document
    the_Report = docx.Document()
    the_Report.add_heading(assessment, 1)
    the_Report = generate_vuln_list(the_Report, assessment, assessment_db, cli_args)

    if ((len(assessment_db) is 0) and cli_args.all_vulns is False):
        exit("["+warn+"]Nothing to report on, quitting...")

    ####################################
    # Write the report in severity order
    ####################################

    logging.info("Writing the critical vulnerability narratives to the report")

    for i in assessment_db:
        if assessment_db[i]['report_rating'] == 'Critical' and cli_args.sC:

            # Grouped Vulnerability Write-up
            if len(assessment_db[i]['vulns']) > 1:
                print '\t['+info+']Multi finding: ', assessment_db[i]['report_title']
                logging.info('['+info+']Multi finding: ' + assessment_db[i]['report_title'])
                the_report = write_multi_vul(assessment_db[i], the_Report, cli_args)

            # Single Vulnerability Write-up
            elif assessment_db[i]['report_rating'] is not None:
                logging.info("["+info+"]" + assessment_db[i]['report_title'] +
                             "(" + assessment_db[i]['report_rating'] + ")")
                print "\t["+info+"]" + assessment_db[i]['report_title'] + " \033[0;0;31m(" + \
                    assessment_db[i]['report_rating'] + ")\033[0m"
                the_Report = write_single_vul(assessment_db[i], the_Report, cli_args)

    logging.info("Writing the high vulnerability narratives to the report")

    for i in assessment_db:
        if assessment_db[i]['report_rating'] == 'High' and cli_args.sH:
            if len(assessment_db[i]['vulns']) > 1:
                logging.info('['+info+']Multi finding: ' + assessment_db[i]['report_title'])
                print '\t['+info+']Multi finding: ', assessment_db[i]['report_title']
                the_report = write_multi_vul(assessment_db[i], the_Report, cli_args)
            elif assessment_db[i]['report_rating'] is not None:
                logging.info("["+info+"]" + assessment_db[i]['report_title'] +
                             "(" + assessment_db[i]['report_rating'] + ")")
                print "\t["+info+"]" + assessment_db[i]['report_title'] + " \033[0;0;35m(" + \
                      assessment_db[i]['report_rating'] + ")\033[0m"
                the_Report = write_single_vul(assessment_db[i], the_Report, cli_args)

    logging.info("Writing the medium vulnerability narratives to the report")

    for i in assessment_db:
        if assessment_db[i]['report_rating'] == 'Medium' and cli_args.sM:
            if len(assessment_db[i]['vulns']) > 1:
                logging.info('['+info+']Multi finding: ' + assessment_db[i]['report_title'])
                print '\t['+info+']Multi finding: ', assessment_db[i]['report_title']
                the_report = write_multi_vul(assessment_db[i], the_Report, cli_args)
            elif assessment_db[i]['report_rating'] is not None:
                logging.info("["+info+"]" + assessment_db[i]['report_title'] +
                             "(" + assessment_db[i]['report_rating'] + ")")
                print "\t["+info+"]" + assessment_db[i]['report_title'] + " \033[0;0;33m(" + \
                      assessment_db[i]['report_rating'] + ")\033[0m"
                the_Report = write_single_vul(assessment_db[i], the_Report, cli_args)

    logging.info("Writing the low vulnerability narratives to the report")

    for i in assessment_db:
        if assessment_db[i]['report_rating'] == 'Low' and cli_args.sL:
            if len(assessment_db[i]['vulns']) > 1:
                logging.info('['+info+']Multi finding: ' + assessment_db[i]['report_title'])
                print '\t['+info+']Multi finding: ', assessment_db[i]['report_title']
                the_report = write_multi_vul(assessment_db[i], the_Report, cli_args)
            elif assessment_db[i]['report_rating'] is not None:
                logging.info("["+info+"]" + assessment_db[i]['report_title'] +
                             "(" + assessment_db[i]['report_rating'] + ")")
                print "\t["+info+"]" + assessment_db[i]['report_title'] + " \033[0;0;34m(" + \
                      assessment_db[i]['report_rating'] + ")\033[0m"
                the_Report = write_single_vul(assessment_db[i], the_Report, cli_args)

    logging.info("Writing the informational vulnerability narratives to the report")

    for i in assessment_db:
        if assessment_db[i]['report_rating'] == 'Informational' and cli_args.sI:
            if len(assessment_db[i]['vulns']) > 1:
                logging.info('['+info+']Multi finding: ' + assessment_db[i]['report_title'])
                print '\t['+info+']Multi finding: ', assessment_db[i]['report_title']
                the_report = write_multi_vul(assessment_db[i], the_Report, cli_args)
            elif assessment_db[i]['report_rating'] is not None:
                logging.info("["+info+"]" + assessment_db[i]['report_title'] +
                             "(" + assessment_db[i]['report_rating'] + ")")
                print "\t["+info+"]" + assessment_db[i]['report_title'] + " \033[0;0;37m(" + \
                      assessment_db[i]['report_rating'] + ")\033[0m"
                the_Report = write_single_vul(assessment_db[i], the_Report, cli_args)

    if cli_args.all_vulns:
        the_Report = write_all_vuln(vuln_ID_to_info_mapping, the_Report, cli_args)
    the_Report = generate_hosts_table(the_Report, assessment)
    save_report(the_Report, assessment)